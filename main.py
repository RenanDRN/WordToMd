import os
import tkinter
from tkinter import filedialog, messagebox
from tkinter import *
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from datetime import datetime
import tkinter.messagebox
import customtkinter
from PIL import Image, ImageTk
import sys
import tldextract

# Função para obter o texto de um parágrafo do documento DOCX, mantendo o estilo de negrito e itálico quando aplicado.
def obter_texto_paragrafo(para):
    """
    Função para obter o texto de um parágrafo do documento DOCX,
    mantendo o estilo de negrito e itálico quando aplicado.

    Args:
        para (docx.text.paragraph.Paragraph): Parágrafo do documento DOCX.

    Returns:
        str: Texto do parágrafo com estilos de formatação aplicados.
    """
    texto = ""
    for run in para.runs:
        run_text = run.text
        if run_text.strip():  # Verifica se o texto não é apenas espaços
            if run.bold and run.italic:
                run_text = f"___{run_text}___"
            elif run.bold:
                run_text = f"**{run_text}**"
            elif run.italic:
                run_text = f"_{run_text}_"
        texto += run_text
    
    return texto

# Função para salvar uma imagem em um diretório específico.
def salvar_imagem(bytes_imagem, diretorio_imagem, nome_imagem, nome_arquivo):
    """
    Função para salvar uma imagem em um diretório específico.

    Args:
        bytes_imagem (bytes): Bytes da imagem a ser salva.
        diretorio_imagem (str): Diretório onde a imagem será salva.
        nome_imagem (str): Nome da imagem a ser salva.

    Returns:
        str: Caminho completo da imagem salva.
    """
    if not os.path.exists(diretorio_imagem):
        os.makedirs(diretorio_imagem)
    nome_imagem = f"{nome_arquivo}_{nome_imagem}"
    caminho_imagem = os.path.join(diretorio_imagem, nome_imagem)
    with open(caminho_imagem, "wb") as arquivo_imagem:
        arquivo_imagem.write(bytes_imagem)
    return ".\\" + caminho_imagem

# Função para identificar e formatar links em um texto.
def formatar_link(texto):
    """
    Função para identificar e formatar links em um texto.

    Args:
        texto (str): Texto contendo URLs.

    Returns:
        str: Texto com URLs formatadas em Markdown.
    """
    # Expressão regular para identificar URLs
    url_pattern = re.compile(r'(http[s]?://[^\s]+)')
    
    def substituir_url(match):
        url = match.group(0)
        dominio = tldextract.extract(url).registered_domain
        return f'[{dominio}]({url})'
    
    # Substituir todas as URLs pelo formato desejado
    texto_formatado = url_pattern.sub(substituir_url, texto)
    return texto_formatado

# Função para converter um documento DOCX para Markdown.
def converter_docx_para_markdown(caminho_docx, caminho_md_saida, diretorio_imagem, nome_arquivo, tipo_cabecalho):
    """
    Função para converter um documento DOCX para Markdown.

    Args:
        caminho_docx (str): Caminho do arquivo DOCX a ser convertido.
        caminho_md_saida (str): Caminho de saída para o arquivo Markdown convertido.
        diretorio_imagem (str): Diretório onde as imagens serão salvas.
    """
    documento = Document(caminho_docx)
    conteudo_markdown = []

    # Adicionando o cabeçalho do arquivo Markdown
    data_atual = datetime.now()
    dia = f'{data_atual.day:02d}'
    mes = f'{data_atual.month:02d}'
    ano = f'{data_atual.year:04d}'

    # Verifica o tipo de cabeçalho selecionado e adiciona ao conteúdo Markdown
    if(tipo_cabecalho == "Topico"):
        conteudo_markdown.append('---')
        conteudo_markdown.append('title: "Titulo da sua documentação"')
        conteudo_markdown.append('type: docs')
        conteudo_markdown.append('last_updated: ')
        conteudo_markdown.append(f'\tdate: "{mes}/{dia}/{ano}"')
        conteudo_markdown.append('\tauthor: "Seu nome"')
        conteudo_markdown.append('sidebar_position: 1')
        conteudo_markdown.append('---\n')
    else:
        conteudo_markdown.append('---')
        conteudo_markdown.append('title: "Titulo da sua documentação"')
        conteudo_markdown.append('type: docs')
        conteudo_markdown.append('menu: ')
        conteudo_markdown.append('\tmain:')
        conteudo_markdown.append('\t\tsidebar_position: 1')
        conteudo_markdown.append('description: "Descrição da sua documentação"')
        conteudo_markdown.append('---\n')

    # Função para adicionar parágrafos ao conteúdo Markdown
    def adicionar_paragrafo(paragrafo):
        texto = obter_texto_paragrafo(paragrafo)
        texto = formatar_link(texto)  # Formatar links no texto do parágrafo
        if paragrafo.style.name.startswith('Heading'):
            nivel = int(re.search(r'\d+', paragrafo.style.name).group())
            conteudo_markdown.append(f"{'#' * nivel} {texto}\n")
        elif paragrafo.style.name == 'Normal' and paragrafo.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
            conteudo_markdown.append(f"<p align='center'>{texto}</p>\n")
        else:
            conteudo_markdown.append(f"{texto}\n")

    # Função para adicionar tabelas ao conteúdo Markdown
    def adicionar_tabela(tabela):
        conteudo_markdown.append("\n")
        # Extrair cabeçalhos da tabela
        cabecalhos = [celula.text.strip() for celula in tabela.rows[0].cells]
        conteudo_markdown.append("| " + " | ".join(cabecalhos) + " |")
        conteudo_markdown.append("| " + " | ".join(["---"] * len(cabecalhos)) + " |")
        
        # Extrair linhas da tabela
        for linha in tabela.rows[1:]:
            dados_linha = []
            for celula in linha.cells:
                texto_celula = obter_texto_paragrafo(celula.paragraphs[0])  # Usar a função para obter o texto formatado
                
                # Verificar se a célula contém uma imagem
                tem_imagem = any(run._element.xpath(".//w:drawing") or run._element.xpath(".//w:pict") for run in celula.paragraphs[0].runs)
                if tem_imagem:
                    for run in celula.paragraphs[0].runs:
                        if run._element.xpath(".//w:drawing") or run._element.xpath(".//w:pict"):
                            for rId in mapa_imagens:
                                if run._element.xpath(f".//*[@r:embed='{rId}']"):
                                    texto_celula += mapa_imagens[rId].strip()  # Concatenar a imagem sem quebra de linha
                                    break
                
                # Verificar se a célula contém um link
                texto_celula = formatar_link(texto_celula)
                
                dados_linha.append(texto_celula)
            
            conteudo_markdown.append("| " + " | ".join(dados_linha) + " |")
        conteudo_markdown.append("\n")                
    contador_imagem = 1
    mapa_imagens = {}

    # Processamento de imagens no documento
    for relacao in documento.part.rels.values():
        if "image" in relacao.reltype:
            parte_imagem = relacao._target
            bytes_imagem = parte_imagem.blob
            nome_imagem = f"imagem{contador_imagem}.png"
            contador_imagem += 1
            caminho_imagem_salva = salvar_imagem(bytes_imagem, diretorio_imagem, nome_imagem, nome_arquivo)
            caminho_imagem_relativa = os.path.relpath(caminho_imagem_salva, os.path.dirname(caminho_md_saida))
            mapa_imagens[relacao.rId] = f"![{nome_imagem}]({caminho_imagem_relativa})\n"

    # Processamento do conteúdo do documento
    for elemento in documento.element.body:
        if isinstance(elemento, CT_P):
            for paragrafo in documento.paragraphs:
                if paragrafo._element == elemento:
                    tem_imagem = any(run._element.xpath(".//w:drawing") or run._element.xpath(".//w:pict") for run in paragrafo.runs)
                    if tem_imagem:
                        for run in paragrafo.runs:
                            if run._element.xpath(".//w:drawing") or run._element.xpath(".//w:pict"):
                                for rId in mapa_imagens:
                                    if run._element.xpath(f".//*[@r:embed='{rId}']"):
                                        conteudo_markdown.append(mapa_imagens[rId])
                                        break
                    else:
                        adicionar_paragrafo(paragrafo)
                    break
        elif isinstance(elemento, CT_Tbl):
            for tabela in documento.tables:
                if tabela._element == elemento:
                    adicionar_tabela(tabela)
                    break

    # Escrever o conteúdo no arquivo Markdown
    with open(caminho_md_saida, "w", encoding="utf-8") as arquivo_md:
        arquivo_md.write("\n".join(conteudo_markdown))

# Função para iniciar a conversão do arquivo DOCX para Markdown
def iniciar_conversao(tipo_cabecalho):
    try:
        # Abre um diálogo para selecionar o arquivo DOCX
        caminho_docx = filedialog.askopenfilename(filetypes=[("Arquivos Word", "*.docx")])
        nome_arquivo = os.path.splitext(os.path.basename(caminho_docx))[0]
        # Abre um diálogo para selecionar o diretório de saída
        diretorio_saida = filedialog.askdirectory()
        caminho_md_saida = os.path.join(diretorio_saida, f"{nome_arquivo}.md")
        diretorio_imagem = os.path.join(diretorio_saida, f"img_{nome_arquivo}")
        # Chama a função de conversão
        converter_docx_para_markdown(caminho_docx, caminho_md_saida, diretorio_imagem, nome_arquivo, tipo_cabecalho)
        messagebox.showinfo("Sucesso", "Conversão realizada com sucesso!")
    except Exception as e:
        messagebox.showerror("Erro", f"Ocorreu um erro: {str(e)}")

def iniciar_conversao_em_lote(tipo_cabecalho):
    try:
        # Abre um diálogo para selecionar o diretório contendo arquivos DOCX
        diretorio_docx = filedialog.askdirectory()
        # Abre um diálogo para selecionar o diretório de saída
        diretorio_saida = filedialog.askdirectory()
        
        for nome_arquivo in os.listdir(diretorio_docx):
            if nome_arquivo.endswith(".docx"):
                caminho_docx = os.path.join(diretorio_docx, nome_arquivo)
                nome_arquivo_sem_extensao = os.path.splitext(nome_arquivo)[0]
                caminho_md_saida = os.path.join(diretorio_saida, f"{nome_arquivo_sem_extensao}.md")
                diretorio_imagem = os.path.join(diretorio_saida, f"img_{nome_arquivo_sem_extensao}")
                # Chama a função de conversão para cada arquivo DOCX
                converter_docx_para_markdown(caminho_docx, caminho_md_saida, diretorio_imagem, nome_arquivo_sem_extensao, tipo_cabecalho)
        
        messagebox.showinfo("Sucesso", "Conversão em lote realizada com sucesso!")
    except Exception as e:
        messagebox.showerror("Erro", f"Ocorreu um erro: {str(e)}")

# Função para obter o caminho absoluto para o recurso, funciona para dev e para PyInstaller
def resource_path(relative_path):
    """ Obtenha o caminho absoluto para o recurso, funciona para dev e para PyInstaller """
    try:
        # PyInstaller cria uma pasta temporária e armazena o caminho nela
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)

# Classe principal da aplicação GUI
class App(customtkinter.CTk):
    def __init__(self):
        super().__init__()

        # Configuração da janela principal
        self.title("WordToMd")
        self.geometry(f"{1100}x{580}")

        # Configuração do layout de grade (2x1)
        self.grid_columnconfigure(1, weight=1)
        self.grid_columnconfigure(2, weight=1)
        self.grid_rowconfigure(0, weight=1)

        # Criação do frame da barra lateral com widgets
        self.sidebar_frame = customtkinter.CTkFrame(self, width=140, corner_radius=0)
        self.sidebar_frame.grid(row=0, column=0, rowspan=1, sticky="nsew")
        self.sidebar_frame.grid_rowconfigure(4, weight=1)

        # Carregar e adicionar a imagem acima do texto "WordToMd"
        imagem_path = resource_path("img/novo-logo-itau-png-sem-fundo.png")  # Use resource_path para obter o caminho correto
        imagem = Image.open(imagem_path)
        imagem = imagem.resize((150, 150), Image.LANCZOS)
        self.logo = customtkinter.CTkImage(light_image=imagem, size=(150, 150))
        self.logo_label = customtkinter.CTkLabel(self.sidebar_frame, image=self.logo, text="")
        self.logo_label.grid(row=0, column=0, padx=20, pady=(20, 10))
        
        self.logo_label_text = customtkinter.CTkLabel(self.sidebar_frame, text="WordToMd", font=customtkinter.CTkFont(size=20, weight="bold"))
        self.logo_label_text.grid(row=1, column=0, padx=20, pady=(10, 10))
        
        self.sidebar_button_1 = customtkinter.CTkButton(self.sidebar_frame, text="Ler Arquivo", command=self.ler_arquivo)
        self.sidebar_button_1.grid(row=2, column=0, padx=20, pady=10)
        self.sidebar_button_2 = customtkinter.CTkButton(self.sidebar_frame, text="Ler Varios Arquivos", command=self.ler_varios_arquivos)
        self.sidebar_button_2.grid(row=3, column=0, padx=20, pady=10)
        self.sidebar_button_3 = customtkinter.CTkButton(self.sidebar_frame, text="Sair", command=self.sair)
        self.sidebar_button_3.grid(row=4, column=0, padx=10, pady=5)
        
        
        self.appearance_mode_label = customtkinter.CTkLabel(self.sidebar_frame, text="Aparência:", anchor="w")
        self.appearance_mode_label.grid(row=5, column=0, padx=20, pady=(10, 0))
        self.appearance_mode_optionemenu = customtkinter.CTkOptionMenu(self.sidebar_frame, values=["Light", "Dark", "System"],
                                                                       command=self.change_appearance_mode_event)
        self.appearance_mode_optionemenu.grid(row=6, column=0, padx=20, pady=(10, 10))
        self.scaling_label = customtkinter.CTkLabel(self.sidebar_frame, text="Escala do aplicativo:", anchor="w")
        self.scaling_label.grid(row=8, column=0, padx=20, pady=(10, 0))
        self.scaling_optionemenu = customtkinter.CTkOptionMenu(self.sidebar_frame, values=["80%", "90%", "100%", "110%", "120%", "130%"],
                                                               command=self.change_scaling_event)
        self.scaling_optionemenu.grid(row=9, column=0, padx=20, pady=(10, 20))

        # Criação do frame para opções de cabeçalho
        self.header_frame = customtkinter.CTkFrame(self)
        self.header_frame.grid(row=0, column=1, padx=(20, 10), pady=(20, 20), sticky="nsew")
        self.header_frame.grid_columnconfigure(0, weight=1)
        self.header_frame.grid_rowconfigure(2, weight=1)

        # Adicionar rótulo de título
        self.options_label = customtkinter.CTkLabel(self.header_frame, text="Selecione o tipo de cabeçalho:", font=customtkinter.CTkFont(size=16, weight="bold"))
        self.options_label.grid(row=0, column=0, padx=20, pady=(20, 10))

        self.optionmenu_1 = customtkinter.CTkOptionMenu(self.header_frame, dynamic_resizing=False,
                                                        values=["Topico", "Sub-Topico"],
                                                        command=self.update_textbox)
        self.optionmenu_1.grid(row=1, column=0, padx=20, pady=(10, 10))

        # Adicionar caixa de texto dinâmica com altura aumentada
        self.dynamic_textbox = customtkinter.CTkTextbox(self.header_frame, state="disabled")
        self.dynamic_textbox.grid(row=2, column=0, padx=20, pady=(10, 10), sticky="nsew")

        # Criação do frame para a caixa de texto principal
        self.textbox_frame = customtkinter.CTkFrame(self)
        self.textbox_frame.grid(row=0, column=2, padx=(10, 20), pady=(20, 20), sticky="nsew")
        self.textbox_frame.grid_columnconfigure(0, weight=1)
        self.textbox_frame.grid_rowconfigure(0, weight=1)

        # Criação da caixa de texto principal
        self.textbox = customtkinter.CTkTextbox(self.textbox_frame)
        self.textbox.grid(row=0, column=0, padx=20, pady=(20, 20), sticky="nsew")

        # Inserir texto de descrição
        self.textbox.insert("0.0", "Descrição:\n\n" + "Autor: Renan dos Reis Negrão\n\n" + "Aplicativo desenvolvido para facilitar a elaboração de novas documentaões técnicas para o GitHub Pages. Através da formatação de um arquivo word a aplicação reconhece e já converte o arquivo para um com extensão .md\n\n" + "Vamos de turma :)\n\n" +  "------------------------------\n\n" +  "Como usar o aplicativo?\n\n" +  "1. Selecione o tipo de cabeçalho desejado!\n\n" + "2. Clique em Ler Arquivo e selecione o arquivo word a ser convertido!\n"+ "OBS: O arquivo Word NÃO pode estar aberto durante a execução do programa!\n\n" + "3. Assim que selecionar o arquivo Word desejado, escolha qual o caminho onde será salvo seu arquivo .MD e a pasta com as imagens (Caso possua no documento)!\n\n")

        # Desabilitar a caixa de texto
        self.textbox.configure(state="disabled")

        # Definir valores padrão
        self.appearance_mode_optionemenu.set("Dark")
        self.scaling_optionemenu.set("100%")

    # Evento para abrir um diálogo de entrada
    def open_input_dialog_event(self):
        dialog = customtkinter.CTkInputDialog(text="Type in a number:", title="CTkInputDialog")
        print("CTkInputDialog:", dialog.get_input())

    # Evento para mudar o modo de aparência
    def change_appearance_mode_event(self, new_appearance_mode: str):
        customtkinter.set_appearance_mode(new_appearance_mode)

    # Evento para mudar a escala do aplicativo
    def change_scaling_event(self, new_scaling: str):
        new_scaling_float = int(new_scaling.replace("%", "")) / 100
        customtkinter.set_widget_scaling(new_scaling_float)

    # Evento para o botão da barra lateral
    def sidebar_button_event(self):
        print("sidebar_button click")

    # Função para ler o arquivo e iniciar a conversão
    def ler_arquivo(self):
        tipo_cabecalho = self.optionmenu_1.get()
        iniciar_conversao(tipo_cabecalho)

    def ler_varios_arquivos(self):
        tipo_cabecalho = self.optionmenu_1.get()
        iniciar_conversao_em_lote(tipo_cabecalho)


    # Função para sair do aplicativo
    def sair(self):
        self.quit()

    # Função para atualizar a caixa de texto dinâmica com base no tipo de cabeçalho selecionado
    def update_textbox(self, value):
        self.dynamic_textbox.configure(state="normal")
        self.dynamic_textbox.delete("1.0", "end")
        conteudo_markdown = []

        # Obter a data atual
        data_atual = datetime.now()
        dia = f'{data_atual.day:02d}'
        mes = f'{data_atual.month:02d}'
        ano = f'{data_atual.year:04d}'

        if value == "Topico":
            conteudo_markdown.append('---')
            conteudo_markdown.append('title: "Titulo da sua documentação"')
            conteudo_markdown.append('type: docs')
            conteudo_markdown.append('last_updated: ')
            conteudo_markdown.append(f'\tdate: "{mes}/{dia}/{ano}"')
            conteudo_markdown.append('\tauthor: "Seu nome"')
            conteudo_markdown.append('sidebar_position: 1')
            conteudo_markdown.append('---\n')
        else:
            conteudo_markdown.append('---')
            conteudo_markdown.append('title: "Titulo da sua documentação"')
            conteudo_markdown.append('type: docs')
            conteudo_markdown.append('menu: ')
            conteudo_markdown.append('\tmain:')
            conteudo_markdown.append('\t\tsidebar_position: 1')
            conteudo_markdown.append('description: "Descrição da sua documentação"')
            conteudo_markdown.append('---\n')

        self.dynamic_textbox.insert("1.0", "\n".join(conteudo_markdown))
        self.dynamic_textbox.configure(state="disabled")

# Inicialização da aplicação
if __name__ == "__main__":
    app = App()
    app.mainloop()